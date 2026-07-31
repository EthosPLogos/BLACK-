"""
Daily job application engine.
Search → Score → Verify → Customize resume → Prepare package → Apply/Queue → Notify.
"""
import json
import subprocess
from datetime import datetime, timezone
from pathlib import Path

from app.config import JOB_APPLY_CONFIG_PATH, JOB_RESUMES_DIR
from app.integrations import job_boards
from app.memory import applications as app_store
from app.services.inference import call_inference

_DEFAULT_CRITERIA: dict = {
    "titles":            [""],
    "location":          "remote",
    "min_score":         60,
    "daily_limit":       10,
    "exclude_companies": [],
    "require_keywords":  [],
    "enabled":           False,
    "run_cron":          "0 9 * * 1-5",
}


# ── Criteria config ────────────────────────────────────────────────────────────

def load_criteria() -> dict:
    if JOB_APPLY_CONFIG_PATH.exists():
        try:
            return {**_DEFAULT_CRITERIA, **json.loads(JOB_APPLY_CONFIG_PATH.read_text())}
        except Exception:
            pass
    return _DEFAULT_CRITERIA.copy()


def save_criteria(criteria: dict) -> None:
    merged = {**_DEFAULT_CRITERIA, **criteria}
    JOB_APPLY_CONFIG_PATH.write_text(json.dumps(merged, indent=2), encoding="utf-8")


def set_enabled(enabled: bool) -> None:
    c = load_criteria()
    c["enabled"] = enabled
    save_criteria(c)


# ── Resume storage ─────────────────────────────────────────────────────────────

def get_base_resume() -> str:
    p = JOB_RESUMES_DIR / "base_resume.txt"
    if p.exists():
        return p.read_text(encoding="utf-8")
    return ""


def save_base_resume(text: str) -> None:
    JOB_RESUMES_DIR.mkdir(parents=True, exist_ok=True)
    (JOB_RESUMES_DIR / "base_resume.txt").write_text(text, encoding="utf-8")


def save_base_resume_bytes(data: bytes, filename: str) -> Path:
    JOB_RESUMES_DIR.mkdir(parents=True, exist_ok=True)
    p = JOB_RESUMES_DIR / f"base_{filename}"
    p.write_bytes(data)
    return p


# ── LLM helpers ────────────────────────────────────────────────────────────────

def _score_job(job: dict, profile: str) -> dict:
    """Ask the LLM to score job fit. Returns {score: int, reason: str}."""
    prompt = (
        f"Score this job for the candidate. Respond with JSON only: "
        f'{{\"score\": 0-100, \"reason\": \"one sentence\"}}\n\n'
        f"JOB TITLE: {job.get('title', '')}\n"
        f"JOB DESCRIPTION:\n{job.get('description') or job.get('snippet', '')}\n\n"
        f"CANDIDATE PROFILE:\n{profile}"
    )
    reply, _ = call_inference(prompt=prompt, system="You score job fit. Respond with JSON only.")
    try:
        # Extract JSON from reply
        import re
        m = re.search(r'\{[^}]+\}', reply)
        if m:
            return json.loads(m.group())
    except Exception:
        pass
    return {"score": 50, "reason": "Could not parse score"}


def _build_custom_resume(job: dict, base_resume: str) -> str:
    """Generate a tailored resume for this specific job."""
    prompt = (
        f"Rewrite this resume tailored specifically for the following job posting.\n\n"
        f"RULES:\n"
        f"- Keep all facts accurate — do not invent skills or experience\n"
        f"- Rewrite the professional summary to target this role directly\n"
        f"- Lead skills section with keywords from the job description\n"
        f"- Reframe experience bullets to use the job posting's language\n"
        f"- Keep length similar to the original\n\n"
        f"JOB TITLE: {job.get('title', '')}\n"
        f"COMPANY: {job.get('company', 'Unknown')}\n"
        f"JOB DESCRIPTION:\n{job.get('description') or job.get('snippet', '')}\n\n"
        f"BASE RESUME:\n{base_resume}"
    )
    reply, _ = call_inference(
        prompt=prompt,
        system="You are an expert resume writer. Output only the tailored resume text. No commentary.",
        tier="frontier",
    )
    return reply


def _build_cover_letter(job: dict, base_resume: str) -> str:
    """Generate a targeted cover letter under 280 words."""
    prompt = (
        f"Write a tailored cover letter for this application. Under 280 words.\n"
        f"- Open directly — no 'I am excited to apply' or similar openers\n"
        f"- Reference the specific role and company\n"
        f"- Connect 2-3 specific experiences from the resume to what the job needs\n"
        f"- End with a clear call to action\n\n"
        f"JOB: {job.get('title', '')} at {job.get('company', 'the company')}\n"
        f"DESCRIPTION:\n{job.get('description') or job.get('snippet', '')}\n\n"
        f"CANDIDATE RESUME:\n{base_resume}"
    )
    reply, _ = call_inference(
        prompt=prompt,
        system="You write concise, direct cover letters. Output only the letter. No commentary.",
        tier="frontier",
    )
    return reply


def _save_application_package(job_id: str, resume_text: str, cover_letter: str) -> Path:
    """Save resume + cover letter to a per-application folder."""
    folder = JOB_RESUMES_DIR / job_id
    folder.mkdir(parents=True, exist_ok=True)
    (folder / "resume.txt").write_text(resume_text, encoding="utf-8")
    (folder / "cover_letter.txt").write_text(cover_letter, encoding="utf-8")

    # Also try to create a DOCX resume if python-docx is available
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        for line in resume_text.split("\n"):
            if line.startswith("# "):
                p = doc.add_paragraph()
                run = p.add_run(line[2:].strip())
                run.bold = True
                run.font.size = Pt(13)
            elif line.strip():
                doc.add_paragraph(line)
        docx_path = folder / "resume.docx"
        doc.save(str(docx_path))
    except Exception:
        pass

    return folder


def _open_url(url: str) -> None:
    """Open job URL in the default browser (macOS)."""
    try:
        subprocess.run(["open", url], check=False, timeout=5)
    except Exception:
        pass


def _notify(title: str, message: str) -> None:
    safe = message.replace('"', "'")[:200]
    subprocess.run(
        ["osascript", "-e",
         f'display notification "{safe}" with title "Mr.Black — {title}" sound name "Ping"'],
        check=False,
    )


# ── Main daily cycle ───────────────────────────────────────────────────────────

def run_apply_cycle(criteria: dict | None = None) -> dict:
    """
    Full daily application cycle.
    Returns a summary dict with counts and details.
    """
    if criteria is None:
        criteria = load_criteria()

    titles = criteria.get("titles") or [""]
    location = criteria.get("location", "remote")
    min_score = criteria.get("min_score", 60)
    daily_limit = criteria.get("daily_limit", 10)
    exclude = {c.lower() for c in criteria.get("exclude_companies", [])}
    base_resume = get_base_resume()

    summary = {
        "run_at": datetime.now(timezone.utc).isoformat(),
        "found": 0,
        "skipped_duplicate": 0,
        "skipped_suspicious": 0,
        "skipped_low_score": 0,
        "applied": 0,
        "queued": 0,
        "errors": [],
        "jobs": [],
    }

    if not base_resume:
        summary["errors"].append("No base resume found — upload one via System → Job Apply before running.")
        return summary

    # 1. Search for jobs
    all_jobs: list[dict] = []
    for title in titles:
        if title.strip():
            found = job_boards.search_jobs(title.strip(), location=location, max_per_source=5)
            all_jobs.extend(found)

    summary["found"] = len(all_jobs)

    # 2. Enrich and verify
    all_jobs = job_boards.enrich_and_verify(all_jobs)

    # 3. Filter and score
    candidates = []
    for job in all_jobs:
        # Deduplicate
        if app_store.exists(job["job_id"]):
            summary["skipped_duplicate"] += 1
            continue

        # Company blacklist
        company_lower = (job.get("company") or "").lower()
        if any(ex in company_lower for ex in exclude if ex):
            app_store.add(job, status=app_store.STATUS_SKIPPED)
            continue

        # Legitimacy check — skip suspicious, but log them for review
        if job["legitimacy"] == "suspicious":
            app_store.add(job, status=app_store.STATUS_SUSPICIOUS)
            summary["skipped_suspicious"] += 1
            summary["jobs"].append({
                "title": job["title"], "company": job.get("company", ""),
                "url": job["url"], "action": "FLAGGED_SUSPICIOUS",
                "flags": job["legitimacy_flags"],
            })
            continue

        # Score fit
        if base_resume:
            scored = _score_job(job, base_resume[:1500])
            job["score"] = scored.get("score", 50)
            job["score_reason"] = scored.get("reason", "")
        else:
            job["score"] = 50
            job["score_reason"] = "No resume for scoring"

        if job["score"] < min_score:
            app_store.add(job, status=app_store.STATUS_SKIPPED)
            summary["skipped_low_score"] += 1
            continue

        candidates.append(job)

    # Sort by score, take top N
    candidates.sort(key=lambda j: j.get("score", 0), reverse=True)
    candidates = candidates[:daily_limit]

    # 4. Build application packages and apply
    for job in candidates:
        try:
            custom_resume = _build_custom_resume(job, base_resume)
            cover_letter = _build_cover_letter(job, base_resume)
            pkg_folder = _save_application_package(job["job_id"], custom_resume, cover_letter)

            # Open the job URL in the browser (user can complete the form with prepared materials)
            _open_url(job["url"])

            # Log as queued — user opens form, materials are ready
            app_store.add(
                job,
                status=app_store.STATUS_QUEUED,
                resume_path=str(pkg_folder),
                cover_letter=cover_letter[:500],
            )
            summary["queued"] += 1
            summary["jobs"].append({
                "title": job["title"],
                "company": job.get("company", "Unknown"),
                "url": job["url"],
                "score": job["score"],
                "action": "QUEUED",
                "package": str(pkg_folder),
                "flags": job.get("legitimacy_flags", []),
            })

        except Exception as exc:
            summary["errors"].append(f"{job.get('title', 'unknown')}: {str(exc)[:100]}")

    # 5. Notify
    total_actioned = summary["queued"] + summary["applied"]
    suspicious_count = summary["skipped_suspicious"]
    msg = (
        f"Found {summary['found']} jobs · {total_actioned} packages ready"
        + (f" · {suspicious_count} suspicious flagged" if suspicious_count else "")
    )
    _notify("Job Apply", msg)

    return summary
